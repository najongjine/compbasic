from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parent
IMG = ROOT / "email_images"
OUT = ROOT / "컴퓨터기초_이메일활용.docx"

# compact_reference_guide preset + named Korean font override
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F6F9"
GREEN = "13795B"
PALE_GREEN = "E8F5EF"
GOLD = "7A5A00"
PALE_GOLD = "FFF4CC"
RED = "9B1C1C"
PALE_RED = "FCE8E6"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "202124"


def set_font(run, size=11, bold=False, color=BLACK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, color="D9E2EC", size="8", side="left"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_font(run, 9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_node = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    run_node.append(text_node)
    fld.append(run_node)
    paragraph._p.append(fld)


def add_para(doc, text="", size=11, bold=False, color=BLACK, align=None,
             before=0, after=6, line=1.25, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size, bold, color, italic)
    return p


def add_callout(doc, label, text, kind="note"):
    colors = {
        "note": (LIGHT_BLUE, DARK_BLUE),
        "tip": (PALE_GOLD, GOLD),
        "safe": (PALE_GREEN, GREEN),
        "warn": (PALE_RED, RED),
    }
    fill, color = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    shade_paragraph(p, fill)
    set_paragraph_border(p, color=color, size="18")
    r1 = p.add_run(f"  {label}  ")
    set_font(r1, 10.5, True, color)
    r2 = p.add_run(text)
    set_font(r2, 10.5, False, BLACK)
    return p


def add_placeholder(doc, text):
    p = add_para(doc, f"[{text} 이미지 첨부]", 10.5, True, RED,
                 WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8, line=1.15)
    shade_paragraph(p, PALE_GOLD)
    set_paragraph_border(p, color=GOLD, size="12")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(item), 10.7, False, BLACK)


def new_decimal_num_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract.findall(qn("w:lvl")):
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
                abstract_id = int(abstract.get(qn("w:abstractNumId")))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        abstract_id = 0
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_steps(doc, items):
    num_id = new_decimal_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.25
        set_font(p.add_run(item), 10.7, False, BLACK)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.24)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.2
        set_font(p.add_run("□ "), 11, True, GREEN)
        set_font(p.add_run(item), 10.7, False, BLACK)


def add_key_table(doc, rows, widths=(2700, 6660), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    if header:
        cells = table.add_row().cells
        for idx, txt in enumerate(header):
            cells[idx].text = ""
            r = cells[idx].paragraphs[0].add_run(txt)
            set_font(r, 10.2, True, WHITE)
            set_cell_shading(cells[idx], BLUE)
    for label, detail in rows:
        cells = table.add_row().cells
        for cell in cells:
            cell.text = ""
        r1 = cells[0].paragraphs[0].add_run(label)
        set_font(r1, 10.2, True, DARK_BLUE)
        set_cell_shading(cells[0], LIGHT_BLUE)
        r2 = cells[1].paragraphs[0].add_run(detail)
        set_font(r2, 10.2, False, BLACK)
    set_table_geometry(table, list(widths), indent=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path, caption, width=6.3, source=None):
    path = Path(path)
    if not path.exists():
        add_placeholder(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    p.paragraph_format.keep_with_next = True
    cap = add_para(doc, f"그림 | {caption}", 9, True, GRAY,
                   WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=1.05)
    if source:
        add_para(doc, f"화면 출처: {source} (캡처: 2026-08-14). 서비스 업데이트에 따라 화면은 달라질 수 있습니다.",
                 8, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=1.05)


def add_unit_title(doc, number, title, subtitle):
    add_para(doc, f"UNIT {number}", 10, True, BLUE, before=0, after=1, line=1.0, keep=True)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(title), 22, True, INK)
    add_para(doc, subtitle, 11, False, GRAY, before=0, after=10, line=1.2)


def add_heading2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), 13, True, BLUE)
    return p


def page_break(doc):
    doc.add_page_break()


def crop_sources():
    jobs = [
        ("03_attachment_help.png", "screen_attachment_help.png", (0, 0, 1265, 1040)),
        ("04_download_help.png", "screen_download_help.png", (0, 0, 1265, 1040)),
        ("05_contacts_help.png", "screen_contacts_help.png", (0, 0, 1265, 1080)),
    ]
    for src_name, dst_name, box in jobs:
        src = IMG / src_name
        dst = IMG / dst_name
        if src.exists():
            im = Image.open(src).convert("RGB")
            im.crop(box).save(dst, quality=94)


crop_sources()
doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

# Styles: exact compact_reference_guide values.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for list_name in ("List Bullet", "List Number"):
    style = doc.styles[list_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10.7)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("컴퓨터 기초 실습 교재  |  이메일 활용"), 8.5, True, GRAY)
footer = section.footer
add_page_number(footer.paragraphs[0])

# Cover - editorial_cover
add_para(doc, "컴퓨터 기초 실습 교재", 11, True, BLUE,
         WD_ALIGN_PARAGRAPH.CENTER, before=94, after=18, line=1.0)
add_para(doc, "이메일 활용", 30, True, INK,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line=1.0)
add_para(doc, "작성·전송부터 첨부파일, 다운로드, 주소록까지", 15, False, DARK_BLUE,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=24, line=1.15)
add_para(doc, "보고 → 따라 하고 → 스스로 확인하는 초보자용 워크북", 11, True, GOLD,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=74, line=1.15)
add_callout(doc, "학습 범위", "이메일 작성과 전송 · 파일 첨부 · 첨부파일 다운로드 · 주소록 관리", "note")
add_para(doc, "학습자 이름  ______________________________", 10.5, False, GRAY,
         WD_ALIGN_PARAGRAPH.CENTER, before=30, after=6)
add_para(doc, "수업 날짜     ______________________________", 10.5, False, GRAY,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=22)
add_para(doc, "기준 환경: Windows 10/11 · 웹브라우저 · Gmail", 9, False, GRAY,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
add_para(doc, "제작일: 2026년 8월 14일", 9, False, GRAY,
         WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
page_break(doc)

# Page 2
add_unit_title(doc, "00", "이 교재를 사용하는 방법", "그림을 먼저 보고 한 단계씩 직접 조작해 보세요.")
add_heading2(doc, "학습 순서")
add_steps(doc, [
    "실제 화면에서 메뉴, 버튼, 입력란의 위치를 찾습니다.",
    "단계별 설명을 읽고 같은 동작을 천천히 따라 합니다.",
    "직접 해보기의 체크 상자에 완료 표시를 합니다.",
    "실수했다면 바로 보내지 말고 받는 사람, 제목, 첨부파일을 다시 확인합니다.",
])
add_callout(doc, "중요", "이 교재의 주소와 사람 이름은 연습용 예시입니다. 실제 수업에서는 본인의 정보와 허락받은 파일만 사용하세요.", "safe")
add_heading2(doc, "차례")
add_key_table(doc, [
    ("UNIT 01", "이메일 작성과 전송"),
    ("UNIT 02", "첨부파일 넣기"),
    ("UNIT 03", "첨부파일 다운로드"),
    ("UNIT 04", "주소록 사용하기"),
    ("마무리", "종합 실습 · 문제 해결 · 확인 문제"),
], widths=(2100, 7260), header=("단원", "학습 내용"))
page_break(doc)

# Page 3
add_unit_title(doc, "준비", "이메일의 기본 구조와 안전", "이메일은 인터넷을 통해 편지와 파일을 주고받는 도구입니다.")
add_heading2(doc, "이메일 주소 읽기")
add_key_table(doc, [
    ("honggildong", "사용자 이름: 누가 사용하는 주소인지 구분합니다."),
    ("@", "골뱅이: 사용자 이름과 서비스 이름을 나눕니다."),
    ("example.com", "도메인: 이메일 서비스를 제공하는 기관이나 사이트입니다."),
], widths=(2400, 6960), header=("부분", "뜻"))
add_heading2(doc, "받은편지함에서 자주 보는 곳")
add_bullets(doc, [
    "받은편지함: 다른 사람이 보낸 메일이 도착합니다.",
    "보낸편지함: 내가 전송한 메일을 다시 확인합니다.",
    "임시보관함: 작성 중이지만 아직 보내지 않은 메일입니다.",
    "스팸함: 광고성 또는 의심스러운 메일이 자동으로 분류됩니다.",
    "휴지통: 삭제한 메일이 일정 기간 보관됩니다.",
])
add_callout(doc, "안전 수칙", "비밀번호, 인증번호, 주민등록번호를 이메일 본문이나 첨부파일로 보내지 마세요. 모르는 사람이 보낸 링크와 파일도 바로 열지 않습니다.", "warn")
page_break(doc)

# Page 4
add_unit_title(doc, "01", "이메일 작성과 전송", "받는 사람, 제목, 본문을 정확히 입력하고 보내기 전에 검토합니다.")
add_figure(doc, IMG / "gmail_ui_10.png", "Google 공식 안내 이미지에서 Gmail 편지쓰기 창 위치 확인", 4.9,
           "Google Gmail 고객센터 - 이메일 작성 및 보내기")
add_heading2(doc, "작성 창의 네 부분")
add_key_table(doc, [
    ("받는 사람", "수신자의 이메일 주소를 정확히 입력합니다."),
    ("제목", "메일의 목적을 짧고 구체적으로 씁니다."),
    ("본문", "인사, 용건, 마무리 인사를 순서대로 씁니다."),
    ("보내기", "모든 내용을 검토한 뒤 마지막에 클릭합니다."),
], widths=(2200, 7160))
page_break(doc)

# Page 5
add_unit_title(doc, "01-1", "편지쓰기 창 열기", "로그인한 Gmail 화면에서 새 메일을 시작합니다.")
add_placeholder(doc, "로그인된 Gmail 받은편지함에서 ‘편지쓰기’ 버튼 위치를 표시한 실제 화면")
add_steps(doc, [
    "웹브라우저에서 mail.google.com을 열고 자신의 계정으로 로그인합니다.",
    "왼쪽 위의 편지쓰기 버튼을 한 번 클릭합니다.",
    "화면 오른쪽 아래에 새 메시지 창이 열리는지 확인합니다.",
    "창이 작으면 오른쪽 위의 전체 화면 버튼으로 크게 엽니다.",
])
add_figure(doc, IMG / "gmail_ui_42.png", "Google 공식 안내 이미지의 편지쓰기와 보내기 버튼", 4.8,
           "Google Gmail 고객센터 - 이메일 작성 및 보내기")
add_callout(doc, "알아두기", "메일을 작성하다 창을 닫아도 대부분 임시보관함에 자동 저장됩니다. 그래도 중요한 내용은 전송 전에 다시 확인하세요.", "tip")
page_break(doc)

# Page 6
add_unit_title(doc, "01-2", "받는 사람과 제목 입력", "주소 한 글자와 제목 한 줄이 메일 전달 결과를 좌우합니다.")
add_heading2(doc, "받는 사람 입력")
add_steps(doc, [
    "받는 사람 칸을 클릭합니다.",
    "상대방의 이메일 주소를 처음부터 끝까지 입력합니다.",
    "자동 완성 목록이 나오면 이름과 주소를 모두 확인한 뒤 선택합니다.",
    "여러 명에게 보낼 때는 각 주소가 작은 이름표 형태로 구분되는지 확인합니다.",
])
add_callout(doc, "참조와 숨은참조", "참조(Cc)는 함께 내용을 알아야 하는 사람, 숨은참조(Bcc)는 다른 수신자에게 주소를 공개하지 않을 때 사용합니다.", "note")
add_heading2(doc, "좋은 제목 만들기")
add_key_table(doc, [
    ("나쁜 예", "안녕하세요 / 질문 / 파일입니다"),
    ("좋은 예", "[컴퓨터 기초] 8월 14일 과제 제출"),
    ("좋은 예", "회의 일정 확인 요청 - 8월 20일 오후 2시"),
], widths=(1800, 7560), header=("구분", "제목"))
add_callout(doc, "전송 전 확인", "받는 사람 주소가 비슷한 동명이인이나 다른 기관 주소가 아닌지 다시 보세요.", "warn")
page_break(doc)

# Page 7
add_unit_title(doc, "01-3", "본문 작성과 예절", "짧고 분명하게 쓰되 필요한 정보는 빠뜨리지 않습니다.")
add_heading2(doc, "기본 순서")
add_key_table(doc, [
    ("1. 인사", "안녕하세요. 컴퓨터 기초 수강생 홍길동입니다."),
    ("2. 용건", "8월 14일 과제 파일을 보내드립니다."),
    ("3. 확인 요청", "첨부파일을 확인해 주시면 감사하겠습니다."),
    ("4. 마무리", "감사합니다. 홍길동 드림"),
], widths=(2100, 7260), header=("구성", "예시 문장"))
add_heading2(doc, "읽기 쉬운 본문")
add_bullets(doc, [
    "한 문단에는 한 가지 내용만 씁니다.",
    "날짜와 시간은 ‘8월 20일(목) 오후 2시’처럼 구체적으로 씁니다.",
    "모두 대문자나 반복된 느낌표 사용을 피합니다.",
    "개인정보와 비밀번호는 본문에 적지 않습니다.",
])
add_callout(doc, "연습 문장", "안녕하세요. 요청하신 교육 일정표를 보내드립니다. 첨부파일을 확인해 주세요. 감사합니다.", "safe")
page_break(doc)

# Page 8
add_unit_title(doc, "01-4", "보내기 전 마지막 점검", "보내기 버튼은 모든 검토가 끝난 뒤 한 번만 클릭합니다.")
add_checklist(doc, [
    "받는 사람의 이름과 이메일 주소가 맞다.",
    "제목만 읽어도 메일의 목적을 알 수 있다.",
    "본문에 인사, 용건, 요청, 마무리가 들어 있다.",
    "날짜, 시간, 장소, 전화번호의 숫자가 맞다.",
    "첨부한다고 쓴 파일이 실제로 첨부되어 있다.",
    "개인정보나 다른 사람에게 공개하면 안 되는 내용이 없다.",
])
add_heading2(doc, "전송 후 확인")
add_steps(doc, [
    "보내기 버튼을 한 번 클릭합니다.",
    "화면 아래쪽의 ‘메일을 보냈습니다’ 안내를 확인합니다.",
    "보낸편지함을 열어 방금 보낸 메일이 있는지 확인합니다.",
    "잘못 보냈다면 즉시 담당자에게 정정 메일을 보냅니다.",
])
add_callout(doc, "주의", "보낸 메일은 상대방 편지함에서 직접 회수할 수 없는 경우가 많습니다. 보내기 전에 확인하는 습관이 가장 중요합니다.", "warn")
page_break(doc)

# Page 9
add_unit_title(doc, "02", "첨부파일 넣기", "문서, 사진, PDF 같은 파일을 이메일과 함께 보냅니다.")
add_figure(doc, IMG / "screen_attachment_help.png", "Gmail 고객센터의 첨부파일 전송 안내 실제 브라우저 화면", 6.3,
           "https://support.google.com/mail/answer/6584?hl=ko")
add_callout(doc, "화면에서 찾기", "‘파일 첨부’ 항목에는 편지쓰기 → 클립 모양 첨부 → 파일 선택 → 열기 순서가 안내되어 있습니다.", "note")
page_break(doc)

# Page 10
add_unit_title(doc, "02-1", "파일 첨부 순서", "파일의 위치와 이름을 먼저 확인하면 실수를 줄일 수 있습니다.")
add_figure(doc, IMG / "gmail_ui_27.png", "Google 공식 안내 이미지에서 클립 모양 첨부 버튼 확인", 5.7,
           "Google Gmail 고객센터 - 이메일 작성 및 보내기")
add_steps(doc, [
    "메일의 받는 사람, 제목, 본문을 먼저 작성합니다.",
    "작성 창 아래쪽의 클립 모양 ‘파일 첨부’ 버튼을 클릭합니다.",
    "파일 선택 창에서 저장된 폴더를 찾아갑니다.",
    "보낼 파일을 한 번 클릭하고 열기 버튼을 클릭합니다.",
    "메일 아래쪽에 파일 이름과 용량이 나타날 때까지 기다립니다.",
])
add_placeholder(doc, "Gmail 작성 창에 파일 이름과 업로드 완료 표시가 나타난 실제 화면")
page_break(doc)

# Page 11
add_unit_title(doc, "02-2", "첨부파일 확인과 안전", "파일을 보내기 전에 이름, 형식, 용량, 내용을 모두 확인합니다.")
add_heading2(doc, "첨부 전 네 가지 확인")
add_key_table(doc, [
    ("파일 이름", "‘새 문서 1’보다 ‘컴퓨터기초_과제_홍길동.docx’처럼 구체적으로 씁니다."),
    ("파일 형식", "문서 .docx, PDF .pdf, 사진 .jpg/.png 등 수신자가 열 수 있는 형식인지 확인합니다."),
    ("파일 용량", "개인 Gmail은 일반적으로 총 25MB까지 첨부하며, 큰 파일은 Drive 링크로 바뀔 수 있습니다."),
    ("파일 내용", "다른 사람의 개인정보, 숨겨진 메모, 잘못된 버전이 들어 있지 않은지 엽니다."),
], widths=(1900, 7460))
add_heading2(doc, "첨부파일 삭제하기")
add_steps(doc, [
    "작성 창 아래의 첨부파일 이름을 찾습니다.",
    "파일 이름 오른쪽의 X 또는 삭제 버튼을 클릭합니다.",
    "잘못된 파일이 사라졌는지 확인한 뒤 올바른 파일을 다시 첨부합니다.",
])
add_callout(doc, "보안", ".exe 같은 실행 파일과 출처를 모르는 압축 파일은 보내거나 열지 않습니다.", "warn")
page_break(doc)

# Page 12
add_unit_title(doc, "02-3", "첨부파일 실습", "전송하지 않고 임시보관함까지 안전하게 연습합니다.")
add_checklist(doc, [
    "바탕 화면이나 연습 폴더에 ‘이메일첨부연습.txt’ 파일을 만들었다.",
    "Gmail에서 새 메시지 창을 열었다.",
    "받는 사람 칸에는 실제 주소 대신 교사가 안내한 연습 주소를 입력했다.",
    "클립 모양 버튼으로 연습 파일을 첨부했다.",
    "파일 이름과 업로드 완료 표시를 확인했다.",
    "전송하지 않고 창을 닫아 임시보관함에 저장했다.",
])
add_callout(doc, "실습 규칙", "교사의 지시가 없으면 실제 주소로 전송하지 않습니다. 첨부 과정을 확인한 뒤 임시보관함에서 삭제해도 됩니다.", "safe")
add_heading2(doc, "실습 기록")
add_para(doc, "첨부한 파일 이름  ______________________________________________", 10.5, color=GRAY)
add_para(doc, "파일이 저장되어 있던 폴더  ___________________________________", 10.5, color=GRAY)
add_para(doc, "어려웠던 단계  _________________________________________________", 10.5, color=GRAY)
page_break(doc)

# Page 13
add_unit_title(doc, "03", "첨부파일 다운로드", "받은 파일을 안전한 폴더에 저장하고 저장 위치를 확인합니다.")
add_figure(doc, IMG / "screen_download_help.png", "Gmail 고객센터의 첨부파일 열기 및 다운로드 안내 실제 브라우저 화면", 6.3,
           "https://support.google.com/mail/answer/30719?hl=ko")
add_placeholder(doc, "받은 Gmail 메시지에서 첨부파일의 다운로드 버튼을 표시한 실제 화면")
page_break(doc)

# Page 14
add_unit_title(doc, "03-1", "안전하게 다운로드하기", "보낸 사람과 파일을 확인한 다음 다운로드합니다.")
add_steps(doc, [
    "받은편지함에서 첨부파일이 있는 메일을 엽니다.",
    "보낸 사람 주소가 알고 있는 사람 또는 기관의 주소인지 확인합니다.",
    "본문 내용과 첨부파일 이름이 자연스럽게 연결되는지 확인합니다.",
    "첨부파일 위에 마우스를 올리고 다운로드 버튼을 클릭합니다.",
    "브라우저의 다운로드 완료 안내에서 폴더 열기 또는 파일 위치 열기를 선택합니다.",
])
add_heading2(doc, "다운로드한 파일 찾기")
add_key_table(doc, [
    ("기본 위치", "대부분 파일 탐색기의 ‘다운로드’ 폴더에 저장됩니다."),
    ("브라우저 확인", "Ctrl+J를 누르면 최근 다운로드 목록을 볼 수 있습니다."),
    ("이름으로 찾기", "파일 탐색기 검색창에 파일 이름의 일부를 입력합니다."),
], widths=(2200, 7160))
add_callout(doc, "안전", "보낸 사람이 익숙해도 예상하지 못한 파일이면 열기 전에 전화나 다른 연락 수단으로 확인하세요.", "warn")
page_break(doc)

# Page 15
add_unit_title(doc, "03-2", "다운로드 파일 정리", "찾기 쉬운 폴더와 알기 쉬운 이름으로 정리합니다.")
add_heading2(doc, "추천 폴더 구조")
add_key_table(doc, [
    ("문서", "컴퓨터기초 > 이메일실습 > 문서"),
    ("사진", "컴퓨터기초 > 이메일실습 > 사진"),
    ("과제", "컴퓨터기초 > 이메일실습 > 제출완료"),
], widths=(1800, 7560), header=("자료 종류", "저장 폴더 예시"))
add_heading2(doc, "파일 이름 바꾸기")
add_steps(doc, [
    "파일 탐색기에서 다운로드한 파일을 한 번 클릭합니다.",
    "F2 키를 누르거나 오른쪽 클릭 후 이름 바꾸기를 선택합니다.",
    "날짜, 내용, 작성자를 포함한 이름을 입력합니다.",
    "Enter를 누르고 확장자(.pdf, .docx 등)가 그대로 있는지 확인합니다.",
])
add_key_table(doc, [
    ("바꾸기 전", "document(3).pdf"),
    ("바꾼 후", "2026-08-14_컴퓨터기초_수업자료.pdf"),
], widths=(1800, 7560))
add_callout(doc, "주의", "파일 이름 끝의 확장자를 임의로 바꾸면 파일이 열리지 않을 수 있습니다.", "tip")
page_break(doc)

# Page 16
add_unit_title(doc, "04", "주소록 사용하기", "이름과 이메일 주소를 저장해 다음 메일에서 빠르게 찾습니다.")
add_figure(doc, IMG / "screen_contacts_help.png", "Google 주소록 고객센터의 연락처 추가 안내 실제 브라우저 화면", 6.3,
           "https://support.google.com/contacts/answer/1069522?hl=ko")
add_placeholder(doc, "로그인된 Google 주소록에서 ‘연락처 만들기’ 입력 창 실제 화면")
page_break(doc)

# Page 17
add_unit_title(doc, "04-1", "연락처 추가와 수정", "이름, 이메일, 전화번호를 정확히 입력하고 저장합니다.")
add_steps(doc, [
    "웹브라우저에서 contacts.google.com을 열고 로그인합니다.",
    "왼쪽 위의 연락처 만들기를 클릭합니다.",
    "이름, 이메일, 필요한 경우 전화번호와 소속을 입력합니다.",
    "이메일 주소의 철자와 @ 뒤의 도메인을 확인합니다.",
    "저장을 클릭하고 연락처 목록에 나타나는지 확인합니다.",
])
add_heading2(doc, "연락처를 사용할 때")
add_bullets(doc, [
    "Gmail 받는 사람 칸에 이름 일부를 입력하면 저장된 주소가 추천됩니다.",
    "동명이인은 이메일 주소와 소속을 함께 확인합니다.",
    "주소가 바뀌면 예전 주소를 그대로 두지 말고 수정합니다.",
    "사용하지 않는 주소는 삭제하거나 ‘이전 주소’라고 메모합니다.",
])
add_callout(doc, "개인정보", "다른 사람의 전화번호와 이메일 주소를 허락 없이 공유하거나 여러 사람에게 공개하지 않습니다.", "warn")
page_break(doc)

# Page 18
add_unit_title(doc, "04-2", "주소록 검색과 그룹", "연락처가 많아지면 검색과 라벨로 정리합니다.")
add_heading2(doc, "연락처 찾기")
add_steps(doc, [
    "Google 주소록 위쪽 검색창을 클릭합니다.",
    "이름, 이메일 주소, 전화번호 중 기억나는 일부를 입력합니다.",
    "검색 결과에서 이름과 이메일 주소를 함께 확인합니다.",
    "연락처를 열어 필요한 정보를 확인하거나 수정합니다.",
])
add_heading2(doc, "라벨로 묶기")
add_key_table(doc, [
    ("가족", "가족 구성원의 연락처"),
    ("컴퓨터 수업", "강사와 같은 반 학습자"),
    ("기관", "은행, 병원, 행정기관 등 공식 연락처"),
], widths=(2200, 7160), header=("라벨 예시", "포함할 연락처"))
add_checklist(doc, [
    "연습용 연락처 한 개를 추가했다.",
    "Gmail 받는 사람 칸에서 이름으로 주소를 불러왔다.",
    "동명이인일 때 주소까지 확인했다.",
    "필요 없는 연습용 연락처를 삭제했다.",
])
page_break(doc)

# Page 19
add_unit_title(doc, "마무리", "이메일 종합 실습", "작성, 첨부, 전송 전 검토, 주소록 저장을 한 번에 연습합니다.")
add_callout(doc, "실습 상황", "컴퓨터 기초 강사에게 수업 소감과 연습 파일을 보내는 상황입니다. 실제 전송은 교사의 허락을 받은 뒤 진행하세요.", "safe")
add_key_table(doc, [
    ("받는 사람", "교사가 안내한 연습용 이메일 주소"),
    ("제목", "[컴퓨터 기초] 이메일 활용 실습 - 홍길동"),
    ("본문", "인사 → 수업 소감 → 첨부파일 안내 → 감사 인사"),
    ("첨부파일", "이메일첨부연습.txt"),
], widths=(2200, 7160), header=("항목", "입력할 내용"))
add_steps(doc, [
    "주소록에서 교사가 안내한 연락처를 찾습니다.",
    "새 메시지 창을 열고 받는 사람과 제목을 입력합니다.",
    "네 문장 이내로 본문을 작성합니다.",
    "연습 파일을 첨부하고 업로드 완료를 확인합니다.",
    "체크리스트로 주소, 제목, 본문, 첨부파일을 검토합니다.",
    "교사의 허락이 있을 때만 보내기 버튼을 클릭합니다.",
])
page_break(doc)

# Page 20
add_unit_title(doc, "마무리-1", "자주 생기는 문제 해결", "오류 메시지와 현재 화면을 먼저 확인합니다.")
add_key_table(doc, [
    ("주소가 빨갛게 보여요", "@와 도메인, 불필요한 띄어쓰기를 확인합니다."),
    ("첨부가 안 돼요", "파일 용량, 파일 형식, 인터넷 연결을 확인합니다."),
    ("파일을 못 찾겠어요", "다운로드 폴더와 Ctrl+J 다운로드 목록을 확인합니다."),
    ("파일이 안 열려요", "확장자와 필요한 프로그램을 확인하고 발신자에게 다시 요청합니다."),
    ("주소가 추천되지 않아요", "주소록 저장 여부와 로그인한 Google 계정을 확인합니다."),
    ("메일을 잘못 보냈어요", "즉시 정정 메일을 보내고 필요한 경우 수신자에게 연락합니다."),
], widths=(2800, 6560), header=("문제", "확인할 내용"))
add_heading2(doc, "도움을 요청할 때 알려 줄 내용")
add_checklist(doc, [
    "어떤 화면에서 문제가 생겼는가?",
    "어떤 버튼을 눌렀는가?",
    "화면에 어떤 오류 문구가 보이는가?",
    "파일 이름과 확장자는 무엇인가?",
    "같은 동작을 다시 하면 문제가 반복되는가?",
])
add_callout(doc, "캡처", "개인정보가 보이지 않도록 가린 뒤 오류 화면을 캡처하면 도움을 받기 쉽습니다.", "tip")
page_break(doc)

# Page 21
add_unit_title(doc, "마무리-2", "확인 문제와 정답", "먼저 답을 가리고 스스로 생각해 보세요.")
questions = [
    "1. 이메일의 목적을 짧게 보여 주는 입력란은 무엇인가요?",
    "2. 파일을 첨부할 때 주로 사용하는 모양은 무엇인가요?",
    "3. 최근 다운로드 목록을 여는 단축키는 무엇인가요?",
    "4. 여러 수신자에게 서로의 주소를 공개하지 않을 때 사용하는 칸은 무엇인가요?",
    "5. 주소록에 저장한 사람을 Gmail에서 빠르게 찾는 방법은 무엇인가요?",
]
for q in questions:
    add_para(doc, q, 10.8, True, DARK_BLUE, before=2, after=3, line=1.15)
    add_para(doc, "답  __________________________________________________________", 10, False, GRAY, before=0, after=7, line=1.0)
page_break(doc)
add_unit_title(doc, "정답", "확인 문제 정답과 마무리", "틀린 문제는 해당 단원으로 돌아가 한 번 더 직접 해보세요.")
add_key_table(doc, [
    ("1번", "제목"),
    ("2번", "클립(종이클립) 모양"),
    ("3번", "Ctrl+J"),
    ("4번", "숨은참조(Bcc)"),
    ("5번", "받는 사람 칸에 저장된 이름의 일부를 입력한다."),
], widths=(1500, 7860))
add_para(doc, "수고하셨습니다!", 20, True, GREEN, WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6, line=1.0)
add_callout(doc, "마지막 습관", "주소·제목·본문·첨부파일을 확인하고 보내는 습관이 안전한 이메일 활용의 핵심입니다.", "safe")
add_para(doc, "화면 자료: Google Gmail 및 Google 주소록 고객센터를 2026-08-14에 직접 캡처했으며, Google 공식 UI 안내 이미지를 함께 사용했습니다.",
         8, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, before=8, after=0, line=1.0)

# Core properties and save.
doc.core_properties.title = "컴퓨터 기초 - 이메일 활용"
doc.core_properties.subject = "이메일 작성·전송, 첨부파일, 다운로드, 주소록 실습 교재"
doc.core_properties.author = "컴퓨터 기초 교재"
doc.core_properties.keywords = "컴퓨터기초, 이메일, Gmail, 첨부파일, 다운로드, 주소록"
doc.save(OUT)
print(OUT)
